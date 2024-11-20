import os
import sys
from io import BytesIO
from datetime import datetime
import pandas as pd
import streamlit as st
from pathlib import Path
from openpyxl.styles import Font, Alignment  # <-- Add Font and Alignment here
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import pyperclip3 as pyperclip

# Load the .env file
load_dotenv()

# Set the page config (this will update the browser tab title)
st.set_page_config(
    page_title="File Categorization",  # Title displayed in the browser tab
    page_icon="logo.png", 
)

# Helper function to get files (excluding folders)
def get_files(path, original_path):
    """
    Retrieves all files in a given path along with their modification time.
    If the modification date is inaccessible due to long path or other issues,
    it still collects the file name and full path.

    Returns:
        list: A list of tuples containing file name, modification time (if available), and full path.
    """
    items = []
    if not os.path.exists(path):
        raise ValueError(f"Base path does not exist: {path}")

    for root, dirs, files in os.walk(path):
        for name in files:
            full_path = os.path.normpath(os.path.join(root, name))
            relative_full_path = full_path.replace(path, original_path, 1)  # Convert to client-relative path
            try:
                # Try to get the modification time
                modified_time = datetime.fromtimestamp(os.path.getmtime(full_path)).strftime('%Y-%m-%d')
            except (OSError, FileNotFoundError):
                # Handle cases where the file is inaccessible or too long
                modified_time = "Unavailable"

            # Append the file name, modification time (or "Unavailable"), and full path
            items.append((name, modified_time, full_path, relative_full_path))

    return items

# Helper function to resolve paths
def resolve_path(path):
    """
    Resolves the path for both client and server environments.
    Handles UNC paths, normalizes separators, and validates accessibility.

    Returns:
        str: Resolved and validated path.

    Raises:
        ValueError: If the path is invalid or inaccessible.
    """
    # Normalize separators to use the correct format for the operating system
    path = os.path.normpath(path)

    # Example of handling mapped drives
    drive_mappings = {
        "Z:": os.getenv('UNC_PATH'),  # Replace with the actual UNC path
        "Y:": r"\\Server\SharedDrive"  # Replace with the actual UNC path for Y:
    }

    # Check if the path starts with a mapped drive and convert to UNC
    drive_letter = path[:2]  # Extract the drive letter, e.g., "Z:"
    if drive_letter in drive_mappings:
        unc_path = os.path.normpath(drive_mappings[drive_letter] + path[2:])
        if not os.path.exists(unc_path):
            raise ValueError(f"Invalid or inaccessible mapped drive path: {unc_path}")
        return unc_path

    # Check for UNC paths (e.g., "\\Server\Share")
    if path.startswith("\\\\"):
        if not os.path.exists(path):
            raise ValueError(f"Invalid or inaccessible UNC path: {path}")
        return path

    # Check for normal local paths
    if not os.path.exists(path):
        raise ValueError(f"Invalid or inaccessible local path: {path}")

    return path

# Word generation in memory
def generate_word(categories):
    doc = Document()

    # Add a table with 2 columns
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # Set column widths
    for cell in table.columns[0].cells:
        cell.width = Inches(3)  # Adjust the width of the first column
    for cell in table.columns[1].cells:
        cell.width = Inches(2)  # Adjust the width of the second column

    # Add headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category / File Name'
    hdr_cells[1].text = 'Last Modified'

    # Set font style for headers
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(12)

    # Add data rows
    for category, items in categories.items():
        if items:
            # Add a row for the category name
            category_row = table.add_row().cells
            category_row[0].text = category
            category_row[1].text = ""
            category_row[0].paragraphs[0].runs[0].font.bold = True

            # Add rows for each file
            for name, modified_time in items:
                file_row = table.add_row().cells
                file_row[0].text = name
                file_row[1].text = modified_time

                # Optionally, set vertical alignment
                for cell in file_row:
                    cell.vertical_alignment = True

    # Save the document to an in-memory buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# PDF generation in memory
def generate_pdf(categories):
    buffer = BytesIO()
    pdf = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    bold_style = styles["Heading4"]
    bold_style.fontSize = 10
    normal_style = styles["BodyText"]
    normal_style.fontSize = 10
    data = [[Paragraph("Category / File Name", bold_style), Paragraph("Last Modified", bold_style)]]

    for category, items in categories.items():
        if items:
            data.append([Paragraph(category, bold_style), ""])
            for name, modified_time in items:
                wrapped_name = Paragraph(name, normal_style)
                data.append([wrapped_name, modified_time])

    table = Table(data, colWidths=[250, 100])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.white),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
    ]))

    pdf.build([table])
    buffer.seek(0)
    return buffer


# Excel generation in memory
def generate_excel(categories):
    data = []
    for category, items in categories.items():
        if items:
            data.append({'Category / File Name': category, 'Last Modified': ''})
            for name, modified_time in items:
                data.append({'Category / File Name': f"   {name}", 'Last Modified': modified_time})

    df = pd.DataFrame(data)
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Files', startrow=1)
        worksheet = writer.sheets['Files']
        headers = ['Category / File Name', 'Last Modified']
        for col_num, value in enumerate(headers, 1):
            cell = worksheet.cell(row=1, column=col_num, value=value)
            cell.font = Font(bold=True, size=12)
            cell.alignment = Alignment(horizontal='left')
        worksheet.column_dimensions['A'].width = 50
        worksheet.column_dimensions['B'].width = 20

        for index, row in df.iterrows():
            worksheet.cell(row=index + 2, column=1, value=row['Category / File Name'])
            worksheet.cell(row=index + 2, column=2, value=row['Last Modified'])

    buffer.seek(0)
    return buffer


# Streamlit UI
st.title("File Categorization Tool")

# Display Button to genrate files on top
directory_path = st.text_input("Enter the directory path:", "")

# Session states for checkboxes and generated files
if 'generate_excel_option' not in st.session_state:
    st.session_state.generate_excel_option = False
if 'generate_word_option' not in st.session_state:
    st.session_state.generate_word_option = False
if 'generate_pdf_option' not in st.session_state:
    st.session_state.generate_pdf_option = False
if 'generated_files' not in st.session_state:
    st.session_state.generated_files = {}
if 'category_selection' not in st.session_state:
    st.session_state.category_selection = {}

# Display checkboxes horizontally using columns
col1, col2, col3 = st.columns(3)

with col1:
    st.session_state.generate_excel_option = st.checkbox("Generate Excel", value=st.session_state.generate_excel_option)
with col2:
    st.session_state.generate_word_option = st.checkbox("Generate Word", value=st.session_state.generate_word_option)
with col3:
    st.session_state.generate_pdf_option = st.checkbox("Generate PDF", value=st.session_state.generate_pdf_option)

# Get files in directory
if directory_path:
    try:
        resolved_path = resolve_path(directory_path)

        if resolved_path and Path(resolved_path).exists():
        # Use both the resolved and client-relative paths
            items = get_files(resolved_path, directory_path)
            
            if items:
                categories = ["CONTRACTUAL", "ARCHITECTURAL", "STRUCTURAL", "SERVICES", "SAFETY", "OTHER"]
                category_selection = {}
                
                                # Generate files if selected
                if st.button("Generate Selected Files"):
                    categorized_data = {cat: [] for cat in categories}
                    for name, (modified_time, category) in st.session_state.category_selection.items():
                        categorized_data[category].append((name, modified_time))

                    # Generate files only if the file type is selected
                    if st.session_state.generate_excel_option:
                        output_excel_buffer = generate_excel(categorized_data)
                        st.session_state.generated_files['excel'] = output_excel_buffer

                    if st.session_state.generate_word_option:
                        output_word_buffer = generate_word(categorized_data)
                        st.session_state.generated_files['word'] = output_word_buffer

                    if st.session_state.generate_pdf_option:
                        output_pdf_buffer = generate_pdf(categorized_data)
                        st.session_state.generated_files['pdf'] = output_pdf_buffer

                # Dynamically show buttons in columns based on selected options
                selected_files = [file for file, selected in {
                    'excel': st.session_state.generate_excel_option,
                    'word': st.session_state.generate_word_option,
                    'pdf': st.session_state.generate_pdf_option
                }.items() if selected]

                # Display download buttons based on the number of selected checkboxes
                num_selected = len(selected_files)

                if num_selected == 1:
                    col1, col2, col3 = st.columns(3)  # Re-initialize columns
                elif num_selected == 2:
                    col1, col2 = st.columns(2)  # Only 2 columns
                else:
                    col1, col2, col3 = st.columns(3)  # All 3 columns

                with col1:
                    if 'excel' in selected_files and 'excel' in st.session_state.generated_files:
                        st.download_button("Download Excel", data=st.session_state.generated_files['excel'], file_name='output.xlsx', mime='application/vnd.ms-excel', key='excel_download_button_1')

                with col2:
                    if 'word' in selected_files and 'word' in st.session_state.generated_files:
                        st.download_button("Download Word", data=st.session_state.generated_files['word'], file_name='output.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document', key='word_download_button_1')

                with col3:
                    if 'pdf' in selected_files and 'pdf' in st.session_state.generated_files:
                        st.download_button("Download PDF", data=st.session_state.generated_files['pdf'], file_name='output.pdf', mime='application/pdf', key='pdf_download_button_1')

                st.write("### Assign Categories")
                for index, item in enumerate(items):
                    name, modified_time, full_path, relative_full_path = item
                    cols = st.columns([3, 1])
                    with cols[0]:
                        if st.button(f"{index+1} {name}", key=f"copy_button_{index}"):
                            try:
                                # pyperclip.copy(full_path)
                                pyperclip.copy(relative_full_path)
                                st.success(f"Path copied to clipboard: {relative_full_path}")
                            except Exception:
                                st.warning(f"Could not copy to clipboard. Please copy manually: {relative_full_path}")
                    with cols[1]:
                        # Track category selections to detect changes
                        category = st.selectbox("Select category", options=categories, key=f"selectbox_{index}", label_visibility="collapsed")
                        category_selection[name] = (modified_time, category)

                # Detect if category selection has changed by comparing with session state
                if category_selection != st.session_state.category_selection:
                    # Update session state with new selection and clear generated files to force regeneration
                    st.session_state.category_selection = category_selection
                    st.session_state.generated_files.clear()

                # Generate files if selected
                if st.button("Generate Selected Files"):
                    categorized_data = {cat: [] for cat in categories}
                    for name, (modified_time, category) in st.session_state.category_selection.items():
                        categorized_data[category].append((name, modified_time))

                    # Generate files only if the file type is selected
                    if st.session_state.generate_excel_option:
                        output_excel_buffer = generate_excel(categorized_data)
                        st.session_state.generated_files['excel'] = output_excel_buffer

                    if st.session_state.generate_word_option:
                        output_word_buffer = generate_word(categorized_data)
                        st.session_state.generated_files['word'] = output_word_buffer

                    if st.session_state.generate_pdf_option:
                        output_pdf_buffer = generate_pdf(categorized_data)
                        st.session_state.generated_files['pdf'] = output_pdf_buffer

                # Dynamically show buttons in columns based on selected options
                selected_files = [file for file, selected in {
                    'excel': st.session_state.generate_excel_option,
                    'word': st.session_state.generate_word_option,
                    'pdf': st.session_state.generate_pdf_option
                }.items() if selected]

                # Display download buttons based on the number of selected checkboxes
                num_selected = len(selected_files)

                if num_selected == 1:
                    col1, col2, col3 = st.columns(3)  # Re-initialize columns
                elif num_selected == 2:
                    col1, col2 = st.columns(2)  # Only 2 columns
                else:
                    col1, col2, col3 = st.columns(3)  # All 3 columns

                with col1:
                    if 'excel' in selected_files and 'excel' in st.session_state.generated_files:
                        st.download_button("Download Excel", data=st.session_state.generated_files['excel'], file_name='output.xlsx', mime='application/vnd.ms-excel', key='excel_download_button_2')

                with col2:
                    if 'word' in selected_files and 'word' in st.session_state.generated_files:
                        st.download_button("Download Word", data=st.session_state.generated_files['word'], file_name='output.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document', key='word_download_button_2')

                with col3:
                    if 'pdf' in selected_files and 'pdf' in st.session_state.generated_files:
                        st.download_button("Download PDF", data=st.session_state.generated_files['pdf'], file_name='output.pdf', mime='application/pdf', key='pdf_download_button_2')


    except ValueError as e:
        st.error(str(e))
